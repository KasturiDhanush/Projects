from docx import Document

# Create a new Document
doc = Document()

# Title Page
doc.add_heading('Multiple Types of Plant Leaf Disease Classification Using Deep Learning', level=1)
doc.add_paragraph('Authors: [Your Name]')
doc.add_paragraph('Institution: [Your Institution]')
doc.add_paragraph('Date: [Submission Date]')
doc.add_page_break()

# Abstract
doc.add_heading('Abstract', level=1)
abstract = (
    "Plant diseases significantly impact agricultural productivity and food security worldwide. "
    "Early and accurate identification of plant leaf diseases is crucial for effective management and control. "
    "Traditional methods of plant disease identification are time-consuming and require expert knowledge. "
    "This paper presents a novel approach for the classification of multiple types of plant leaf diseases using transfer learning. "
    "Leveraging pre-trained deep learning models, we fine-tuned these models on a large dataset of plant leaf images representing various diseases and healthy leaves. "
    "Our approach utilizes state-of-the-art convolutional neural networks (CNNs) such as VGG19 and ResNet pre-trained on the ImageNet dataset. "
    "These models were retrained using a dataset of labeled plant leaf images to capture specific disease features. "
    "The transfer learning technique allows the model to benefit from the pre-trained weights, significantly reducing the need for extensive computational resources and large datasets for training from scratch. "
    "Experimental results demonstrate that our transfer learning-based models achieve high accuracy in classifying multiple types of plant leaf diseases, outperforming traditional machine learning methods. "
    "The ResNet model, in particular, showed the highest accuracy, suggesting its robustness in feature extraction for plant disease recognition. "
    "The proposed method can be deployed as a practical tool for farmers and agricultural professionals, enabling rapid and precise disease detection which is essential for timely intervention and reducing crop losses."
)
doc.add_paragraph(abstract)
doc.add_page_break()

# Table of Contents
doc.add_heading('Table of Contents', level=1)
doc.add_paragraph('1. Abstract')
doc.add_paragraph('2. Introduction')
doc.add_paragraph('3. Literature Review')
doc.add_paragraph('4. Methodology')
doc.add_paragraph('5. Results and Discussion')
doc.add_paragraph('6. Conclusion')
doc.add_paragraph('7. References')
doc.add_page_break()

# Introduction
doc.add_heading('Introduction', level=1)
introduction = (
    "Agriculture plays a pivotal role in the global economy, providing food, raw materials, and employment opportunities. "
    "Ensuring the health of crops is paramount for sustaining agricultural productivity. However, plant diseases pose significant challenges, "
    "causing substantial losses in yield and quality. Early and accurate detection of plant diseases is crucial for effective management and control. "
    "Traditional methods of disease detection, which rely on manual inspection by experts, are time-consuming, labor-intensive, and often subjective. "
)
doc.add_paragraph(introduction)

# Save the document
doc_path = "Plant_Leaf_Disease_Classification_Project.docx"
doc.save(doc_path)

print(f"Document saved at {doc_path}")
